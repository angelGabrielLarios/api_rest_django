from django.shortcuts import render

# Create your views here.
from rest_framework.decorators import api_view
from rest_framework.response import Response

from django.http import FileResponse, Http404, HttpResponse

from.models import ModelVideoIndexer
from dotenv import load_dotenv


from docx import Document
from docx.shared import Inches
from datetime import datetime
from .videoindexer.code import get_video_transcript
from .openia_azure.code import generate_minuta_text
from .generate_word.code import generate_word
from django.http import JsonResponse


from rest_framework import generics
from rest_framework import status
from .models import Video
from .serializers import VideoSerializer


import os

load_dotenv()


@api_view(['GET'])
def generate_minutes_of_video_meeting(request):
    try:
        prompt_content = get_video_transcript()
        chat_generated_minute = generate_minuta_text(prompt_content)
        file_path = generate_word(chat_generated_minute)

        # Imprimir la ruta del archivo para depuración
        print(f"File path: {file_path}")

        # Asegúrate de que el archivo existe antes de intentar devolverlo
        if not os.path.exists(file_path):
            raise Http404('File not found.')

        # Verificar permisos de archivo
        if not os.access(file_path, os.R_OK):
            raise PermissionError('File is not readable.')

        # Leer el contenido del archivo
        with open(file_path, 'r') as file:
            file_content = file.read()

        # Devolver el contenido del archivo como respuesta HTTP
        response = HttpResponse(file_content, content_type='text/markdown')
        response['Content-Disposition'] = f'attachment; filename="minuta.md"'
        return response

        
    except Exception as e:
        # Si ocurre una excepción, se captura aquí
        mensaje_error = "Ocurrió un error en la vista: {}".format(str(e))
        return JsonResponse({"error": mensaje_error}, status=500)



# views.py



class VideoUploadView(generics.CreateAPIView):
    # Configuración de la vista para manejar la subida de archivos
    queryset = Video.objects.all()
    serializer_class = VideoSerializer
    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        
        # Construir la respuesta con información adicional
        file_url = serializer.data['file']
        file_name = request.data['file'].name
        file_size = request.data['file'].size
        mime_type = request.data['file'].content_type
        
        response_data = {
            "message": "File uploaded successfully.",
            "file_url": file_url,
            "file_name": file_name,
            "file_size": f"{file_size / (1024 * 1024):.2f} MB",  # Convertir bytes a MB
            "mime_type": mime_type
        }
        
        headers = self.get_success_headers(serializer.data)
        return Response(response_data, status=status.HTTP_201_CREATED, headers=headers)



@api_view(['GET'])
def download_file(request, filename):
    base_dir = "C:\\Users\\angelgabriel\\Documents\\casystems_empresa\\proyects\\api_rest_django_request_chatgpt\\myproject"
    file_path = os.path.join(base_dir, filename)
    
    if os.path.exists(file_path):
        return FileResponse(open(file_path, 'rb'), as_attachment=True, filename=filename)
    else:
        raise Http404("El archivo no existe")
    

def generar_minuta(request):
    # Aquí obtienes la transcripción de tu fuente, por ejemplo, Azure Video Indexer
    transcripcion = "Aquí va la transcripción del video."

    # Obtener la minuta generada por ChatGPT
    minuta_generada = obtener_minuta(transcripcion)

    # Crear un nuevo documento de Word
    doc = Document()
    doc.add_heading('Minuta de Reunión', 0)

    # Aquí puedes agregar lógica para dividir la minuta en secciones y tablas

    # Ejemplo de añadir la minuta completa como texto
    doc.add_paragraph(minuta_generada)

    # Guardar el documento en un archivo temporal
    doc_path = 'minuta.docx'
    doc.save(doc_path)

    # Enviar el archivo como respuesta
    with open(doc_path, 'rb') as f:
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={os.path.basename(doc_path)}'
        return response

def generar_minuta1(request):
    # Datos de ejemplo
    fecha_minuta = datetime.now().strftime("%Y-%m-%d")
    plataforma = "Zoom"  # Esto debería venir de la solicitud
    participantes = [
        {"nombre": "Juan Pérez", "correo": "juan@example.com", "firma": ""},
        {"nombre": "María López", "correo": "maria@example.com", "firma": ""}
    ]
    temas = [
        {"tema": "Introducción", "descripcion": "Presentación del proyecto"},
        {"tema": "Planificación", "descripcion": "Definición de las tareas"}
    ]
    acuerdos = [
        {"acuerdo": "Realizar informe semanal", "responsable": "Juan Pérez"}
    ]

    # Crear un nuevo documento de Word
    doc = Document()

    # Agregar un título
    doc.add_heading('Minuta de Reunión', 0)

    # Agregar la fecha y la plataforma
    doc.add_paragraph(f'Fecha: {fecha_minuta}')
    doc.add_paragraph(f'Plataforma: {plataforma}')

    # Agregar la tabla de participantes
    doc.add_heading('Participantes', level=1)
    tabla_participantes = doc.add_table(rows=1, cols=3)
    hdr_cells = tabla_participantes.rows[0].cells
    hdr_cells[0].text = 'Nombre'
    hdr_cells[1].text = 'Correo'
    hdr_cells[2].text = 'Firma'
    for participante in participantes:
        row_cells = tabla_participantes.add_row().cells
        row_cells[0].text = participante["nombre"]
        row_cells[1].text = participante["correo"]
        row_cells[2].text = participante["firma"]

    # Agregar la tabla de temas a tratar
    doc.add_heading('Temas a Tratar', level=1)
    tabla_temas = doc.add_table(rows=1, cols=2)
    hdr_cells = tabla_temas.rows[0].cells
    hdr_cells[0].text = 'Tema'
    hdr_cells[1].text = 'Descripción'
    for tema in temas:
        row_cells = tabla_temas.add_row().cells
        row_cells[0].text = tema["tema"]
        row_cells[1].text = tema["descripcion"]

    # Agregar la tabla de acuerdos
    doc.add_heading('Acuerdos', level=1)
    tabla_acuerdos = doc.add_table(rows=1, cols=2)
    hdr_cells = tabla_acuerdos.rows[0].cells
    hdr_cells[0].text = 'Acuerdo'
    hdr_cells[1].text = 'Responsable'
    for acuerdo in acuerdos:
        row_cells = tabla_acuerdos.add_row().cells
        row_cells[0].text = acuerdo["acuerdo"]
        row_cells[1].text = acuerdo["responsable"]

    # Guardar el documento en un archivo temporal
    doc_path = 'minuta.docx'
    doc.save(doc_path)

    # Enviar el archivo como respuesta
    with open(doc_path, 'rb') as f:
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={os.path.basename(doc_path)}'
        return response